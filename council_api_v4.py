from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import asyncio
from playwright.async_api import async_playwright
import time
from pathlib import Path
from docx import Document
import threading

app = Flask(__name__)
CORS(app)

ALL_AGENTS = {
    "Analyst": {
        "url": "https://www.genspark.ai/agents?id=59557c0c-1493-4814-8de1-b304a02665ba",
        "name": "The Analyst"
    },
    "Strategist": {
        "url": "https://www.genspark.ai/agents?id=ec8b3f1d-80cc-4d1e-af97-a29c09038c5b",
        "name": "The Strategist"
    },
    "DevilsAdvocate": {
        "url": "https://www.genspark.ai/agents?id=fa0f916e-2555-406d-85c8-883723568885",
        "name": "The Devil's Advocate"
    },
    "Creative": {
        "url": "https://www.genspark.ai/agents?id=59704dd9-9275-4e1a-816a-387d39614dc0",
        "name": "The Creative"
    },
    "FinancialAnalyst": {
        "url": "https://www.genspark.ai/agents?id=04d0e973-94e4-433f-ba6e-ea7a320279fa",
        "name": "The Financial Analyst"
    },
    "Synthesiser": {
        "url": "https://www.genspark.ai/agents?id=ba6db65e-743e-4728-8f70-8bfdc7c18056",
        "name": "The Synthesiser"
    }
}

COUNCIL_PRESETS = {
    "full": ["Analyst", "Strategist", "DevilsAdvocate", "Creative", "FinancialAnalyst"],
    "core": ["Analyst", "Strategist", "DevilsAdvocate", "Creative"],
    "quick": ["Analyst", "Strategist"]
}

active_sessions = {}
session_counter = int(time.time())

def create_word_document(session_id, question, responses, synthesis, doc_type="full"):
    doc = Document()
    doc.add_heading('AI COUNCIL DELIBERATION', 0)
    doc.add_paragraph(f"Session: {session_id}")
    doc.add_paragraph(f"Question: {question}")
    doc.add_paragraph()
    
    if doc_type == "executive":
        doc.add_heading('EXECUTIVE SUMMARY', 1)
        doc.add_paragraph(synthesis)
    else:
        doc.add_heading('SYNTHESIS', 1)
        doc.add_paragraph(synthesis)
        doc.add_paragraph()
        doc.add_heading('ADVISOR PERSPECTIVES', 1)
        for advisor, response in responses.items():
            doc.add_heading(ALL_AGENTS[advisor]["name"], 2)
            doc.add_paragraph(response)
    
    filename = f"Council_{doc_type}_{session_id}.docx"
    filepath = Path("/tmp") / filename
    doc.save(str(filepath))
    return str(filepath)

async def consult_advisor(page, advisor_name, question):
    agent_url = ALL_AGENTS[advisor_name]["url"]
    agent_name = ALL_AGENTS[advisor_name]["name"]
    
    print(f"📞 Consulting {agent_name}...")
    
    try:
        await page.goto(agent_url, timeout=180000, wait_until="networkidle")
        
        # Wait extra time for page to fully render
        await asyncio.sleep(5)
        
        # Try multiple selectors with retries
        selectors = [
            'textarea.active',
            'textarea.search-input',
            'textarea[name="query"]',
            'textarea[placeholder*="Ask"]',
            'textarea',
            'input[type="text"]'
        ]
        
        input_field = None
        for attempt in range(3):
            for selector in selectors:
                try:
                    input_field = await page.wait_for_selector(selector, timeout=10000, state="visible")
                    if input_field:
                        # Verify it's actually interactable
                        is_visible = await input_field.is_visible()
                        is_enabled = await input_field.is_enabled()
                        if is_visible and is_enabled:
                            print(f"✅ Found input field with selector: {selector}")
                            break
                        else:
                            input_field = None
                except:
                    continue
            
            if input_field:
                break
            
            print(f"⏳ Retry {attempt + 1}/3 - waiting for input field...")
            await asyncio.sleep(3)
        
        if not input_field:
            print(f"❌ Could not find input field for {agent_name}")
            return f"[Could not submit to {agent_name}]"
        
        # Submit question
        await input_field.click()
        await asyncio.sleep(1)
        await input_field.fill(question)
        await asyncio.sleep(1)
        await page.keyboard.press("Enter")
        
        print(f"✅ Submitted to {agent_name}")
        print(f"⏳ Waiting for {agent_name} response (180s)...")
        
        await asyncio.sleep(180)
        
        body_text = await page.evaluate("document.body.innerText")
        lines = [line.strip() for line in body_text.split('\n') if len(line.strip()) > 50]
        response = '\n'.join(lines[-20:])
        
        print(f"✅ {agent_name} response captured ({len(response)} chars)")
        return response
        
    except Exception as e:
        print(f"❌ Error consulting {agent_name}: {str(e)}")
        return f"[Error: {str(e)}]"

async def run_council_session(session_id, question, context, selected_advisors):
    print(f"\n🏛️ STARTING COUNCIL SESSION {session_id}")
    
    active_sessions[session_id]["status"] = "in_progress"
    active_sessions[session_id]["progress"] = "Initializing..."
    
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=['--no-sandbox', '--disable-dev-shm-usage', '--disable-blink-features=AutomationControlled']
        )
        
        context_obj = await browser.new_context(
            viewport={'width': 1920, 'height': 1080},
            user_agent='Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
        )
        
        active_sessions[session_id]["progress"] = "Consulting advisors..."
        
        tasks = []
        for advisor in selected_advisors:
            page = await context_obj.new_page()
            tasks.append(consult_advisor(page, advisor, question))
        
        responses = await asyncio.gather(*tasks)
        
        advisor_responses = {}
        for i, advisor in enumerate(selected_advisors):
            advisor_responses[advisor] = responses[i]
        
        active_sessions[session_id]["responses"] = advisor_responses
        active_sessions[session_id]["progress"] = "Synthesizing..."
        
        synthesis_page = await context_obj.new_page()
        synthesis_prompt = f"Question: {question}\n\n"
        for advisor, response in advisor_responses.items():
            synthesis_prompt += f"{ALL_AGENTS[advisor]['name']}: {response}\n\n"
        synthesis_prompt += "Synthesize into a unified recommendation."
        
        try:
            await synthesis_page.goto(ALL_AGENTS["Synthesiser"]["url"], timeout=180000, wait_until="networkidle")
            await asyncio.sleep(5)
            
            selectors = ['textarea.active', 'textarea.search-input', 'textarea[name="query"]', 'textarea']
            
            synth_input = None
            for selector in selectors:
                try:
                    synth_input = await synthesis_page.wait_for_selector(selector, timeout=10000, state="visible")
                    if synth_input and await synth_input.is_visible():
                        break
                except:
                    continue
            
            if synth_input:
                await synth_input.click()
                await asyncio.sleep(1)
                await synth_input.fill(synthesis_prompt)
                await asyncio.sleep(1)
                await synthesis_page.keyboard.press("Enter")
                await asyncio.sleep(180)
            
            body_text = await synthesis_page.evaluate("document.body.innerText")
            lines = [line.strip() for line in body_text.split('\n') if len(line.strip()) > 50]
            synthesis = '\n'.join(lines[-30:])
            
        except Exception as e:
            synthesis = f"[Synthesis error: {str(e)}]"
        
        active_sessions[session_id]["synthesis"] = synthesis
        
        await browser.close()
        
        full_doc = create_word_document(session_id, question, advisor_responses, synthesis, "full")
        exec_doc = create_word_document(session_id, question, advisor_responses, synthesis, "executive")
        
        active_sessions[session_id]["full_report_path"] = full_doc
        active_sessions[session_id]["executive_report_path"] = exec_doc
        active_sessions[session_id]["status"] = "complete"
        active_sessions[session_id]["progress"] = "Complete!"
        
        print(f"✅ COUNCIL SESSION {session_id} COMPLETE!")

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "agents": 6, "version": "4.0"})

@app.route('/api/council/start', methods=['POST'])
def start_council():
    global session_counter
    
    data = request.json
    question = data.get('question')
    context = data.get('context', '')
    preset = data.get('preset', 'core')
    
    if not question:
        return jsonify({"error": "Question required"}), 400
    
    selected_advisors = COUNCIL_PRESETS.get(preset, COUNCIL_PRESETS['core'])
    
    session_counter += 1
    session_id = str(session_counter)
    
    active_sessions[session_id] = {
        "question": question,
        "status": "started",
        "progress": "Starting...",
        "responses": {},
        "synthesis": ""
    }
    
    def run_async():
        asyncio.run(run_council_session(session_id, question, context, selected_advisors))
    
    thread = threading.Thread(target=run_async)
    thread.start()
    
    return jsonify({"session_id": session_id, "status": "started"})

@app.route('/api/council/status/<session_id>', methods=['GET'])
def get_status(session_id):
    session = active_sessions.get(session_id)
    if not session:
        return jsonify({"error": "Session not found"}), 404
    
    return jsonify({
        "status": session["status"],
        "progress": session["progress"]
    })

@app.route('/api/council/download/full/<session_id>', methods=['GET'])
def download_full(session_id):
    session = active_sessions.get(session_id)
    if not session or session["status"] != "complete":
        return jsonify({"error": "Report not ready"}), 404
    
    filepath = session.get("full_report_path")
    if not filepath or not Path(filepath).exists():
        return jsonify({"error": "File not found"}), 404
    
    return send_file(filepath, as_attachment=True)

@app.route('/api/council/download/executive/<session_id>', methods=['GET'])
def download_executive(session_id):
    session = active_sessions.get(session_id)
    if not session or session["status"] != "complete":
        return jsonify({"error": "Report not ready"}), 404
    
    filepath = session.get("executive_report_path")
    if not filepath or not Path(filepath).exists():
        return jsonify({"error": "File not found"}), 404
    
    return send_file(filepath, as_attachment=True)

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    
    print("\n" + "="*80)
    print("🏛️  AI COUNCIL API v4.0")
    print("="*80)
    print(f"\n✅ Server: http://0.0.0.0:{port}")
    print("✅ 6 Council Members Ready")
    print("\nPress Ctrl+C to stop\n")
    
    app.run(host='0.0.0.0', port=port, debug=False)
